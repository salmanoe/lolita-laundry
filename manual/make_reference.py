# Membrand reference.docx Pandoc: warna heading navy/ocean, font, logo di header.
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY  = RGBColor(0x00, 0x2F, 0x6C)
OCEAN = RGBColor(0x03, 0x69, 0xA1)
FONT  = "Segoe UI"

doc = docx.Document("reference.docx")

def style(name, color=None, size=None, bold=True, font=FONT):
    try:
        s = doc.styles[name]
    except KeyError:
        return
    f = s.font
    if color is not None: f.color.rgb = color
    if size  is not None: f.size = Pt(size)
    f.bold = bold
    f.name = font

style("Title",     NAVY,  26)
style("Subtitle",  OCEAN, 13, bold=False)
style("Heading 1", NAVY,  18)
style("Heading 2", OCEAN, 14)
style("Heading 3", NAVY,  12.5)
style("Heading 4", OCEAN, 11.5)

# Body font default
try:
    n = doc.styles["Normal"].font
    n.name = "Calibri"; n.size = Pt(11)
except KeyError:
    pass

# Logo di header (kanan atas), tampil di semua halaman.
sec = doc.sections[0]
sec.header.is_linked_to_previous = False
p = sec.header.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run().add_picture("img/logo.png", height=Inches(0.38))

doc.save("reference.docx")
print("reference.docx berbrand: Title/Heading diwarnai, logo di header.")
